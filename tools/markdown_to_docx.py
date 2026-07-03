from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_font(style, name: str, size: float | None = None, color: str = "000000") -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    style.font.color.rgb = None
    rpr = style._element.get_or_add_rPr()
    color_el = rpr.find(qn("w:color"))
    if color_el is None:
        color_el = OxmlElement("w:color")
        rpr.append(color_el)
    color_el.set(qn("w:val"), color)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_font(normal, "Arial", 11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        style = doc.styles[name]
        set_font(style, "Arial", size, color)
        style.font.bold = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        set_font(style, "Arial", 11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    set_font(code, "Consolas", 9, "202124")
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0


def add_title(doc: Document, title: str, subtitle: str | None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(26)
    run.font.bold = False

    if subtitle:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(16)
        run = sub.add_run(subtitle)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(11)
        run.font.color.rgb = None


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(set(cell) <= {"-", ":", " "} for cell in rows[1]):
        return rows[0], rows[2:]
    return rows[0], rows[1:]


def add_table(doc: Document, table_lines: list[str]) -> None:
    header, body = parse_table(table_lines)
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = True
    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in body:
        cells = table.add_row().cells
        for idx, text in enumerate(row[: len(cells)]):
            cells[idx].text = text
    doc.add_paragraph()


def convert_markdown(markdown: str, out_path: Path, source_name: str) -> None:
    doc = Document()
    configure_styles(doc)

    lines = markdown.splitlines()
    title = None
    body_start = 0
    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        body_start = 1
    add_title(doc, title or source_name, source_name)

    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_table(doc, table_lines)
            table_lines = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            paragraph = doc.add_paragraph(style="Code Block")
            for idx, line in enumerate(code_lines):
                if idx:
                    paragraph.add_run().add_break(WD_BREAK.LINE)
                paragraph.add_run(line)
            code_lines = []

    for raw in lines[body_start:]:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("### "):
            add_inline_runs(doc.add_paragraph(style="Heading 3"), line[4:].strip())
        elif line.startswith("## "):
            add_inline_runs(doc.add_paragraph(style="Heading 2"), line[3:].strip())
        elif line.startswith("# "):
            add_inline_runs(doc.add_paragraph(style="Heading 1"), line[2:].strip())
        elif re.match(r"^\d+\.\s+", line):
            add_inline_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            add_inline_runs(doc.add_paragraph(style="List Bullet"), line[2:])
        elif line == "---":
            doc.add_paragraph()
        else:
            add_inline_runs(doc.add_paragraph(), line)

    flush_table()
    flush_code()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--source-name", default=None)
    args = parser.parse_args()
    convert_markdown(
        args.input.read_text(encoding="utf-8"),
        args.output,
        args.source_name or args.input.as_posix(),
    )


if __name__ == "__main__":
    main()
